import io
import tempfile
from docx import Document
from docx.shared import Inches
from datetime import datetime
import win32print
import win32api
from tkinter import Tk, Button, messagebox
from barcode import EAN13
from barcode.writer import ImageWriter
from PIL import Image

def generate_barcode_png(barcode_text):
    ean = EAN13(barcode_text, writer=ImageWriter())
    with io.BytesIO() as image_stream:
        ean.write(image_stream, options={"write_text": False})
        image_stream.seek(0)
        barcode_image = Image.open(image_stream)
        return barcode_image

def create_label_file(label_info):
    # Create a new Word document
    doc = Document()

    # Add text to the document
    doc.add_heading("ATMART MALL", level=1)
    doc.add_paragraph("Item Name    : " + label_info["Item Name"])
    doc.add_paragraph("Weight       : " + str(label_info["Weight"]))  # Convert to string
    doc.add_paragraph("Price per Kg : " + label_info["Price per Kg"])
    doc.add_paragraph("Total Price  : " + label_info["Total Price"])
    doc.add_paragraph("Date & Time  : " + label_info["Date & Time"])
    doc.add_paragraph("THANK YOU!")

    # Save the barcode PNG to a temporary file
    barcode_image = label_info["Barcode Image"]
    barcode_image_path = tempfile.mktemp(suffix=".png")
    barcode_image.save(barcode_image_path)
    # Add the barcode image to the document
    doc.add_picture(barcode_image_path, width=Inches(2))  # Adjust width as needed

    # Save the document
    file_path = tempfile.mktemp(suffix=".docx")
    doc.save(file_path)

    return file_path

def print_file(printer_name, file_path):
    try:
        win32api.ShellExecute(0, "print", file_path, '/d:"%s"' % printer_name, ".", 0)
    except Exception as e:
        messagebox.showerror("Printing Error", f"Failed to print: {e}")

def on_print_button_click():
    label_info = {
        "Item Name": "Sample Item",
        "Weight": 2.5,
        "Price per Kg": "$10.00",
        "Total Price": "$25.00",
        "Barcode": "123456789012",
        "Date & Time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    # Generate barcode image in PNG format
    barcode_image = generate_barcode_png(label_info["Barcode"])
    label_info["Barcode Image"] = barcode_image

    # Create label file
    file_path = create_label_file(label_info)
    print_file("POS80 Printer", file_path)

# GUI setup
root = Tk()
root.title("Label Printer")

print_button = Button(root, text="Print Label", command=on_print_button_click)
print_button.pack(pady=200,padx=200)

root.mainloop()
