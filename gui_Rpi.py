import tkinter as tk
from tkinter import ttk, messagebox,simpledialog
from PIL import Image, ImageTk
# import barcode 
# from barcode.writer import ImageWriter
from datetime import datetime
import threading
import cv2
from ultralytics import YOLO
import numpy as np
import mysql.connector as mysql
import tempfile
# import win32print
import serial
import time
from docx import Document
from docx.shared import Inches
from decimal import Decimal
import io



# Global variables
running = False
video_thread = None
model = YOLO(r'/home/asus/Desktop/Interface/guienv/lib/python3.11/site-packages/Model1.pt') 
class_names = model.names
price = {"apple": 50, "banana": 40, "cherry": 60, "grapes": 70, "kiwi": 80, "lemon": 90, "mango": 100, "orange": 110, "papaya": 120, "pear": 130, "pineapple": 140, "plum": 150, "pomegranate": 160, "strawberry": 170, "watermelon": 180,"corn": 190}
weight=2.00
item_detected = "apple"
flag=0
global update_id

 # Define the serial port and the baud rate
ser = serial.Serial('COM5', 115200)  # Replace 'COM3' with your Arduino's serial port
    # Give some time for the serial connection to initialize
time.sleep(2)


def create_label_text(item_name, weight, price_per_kg, total_price):

     # Get current date and time
    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # Format the label text
    label_text = (
        "\n----------------------------------------\n"
        "               ATMART MALL              \n"
        "----------------------------------------\n"
        f"Item Name    : {item_name}\n"
        f"Weight       : {weight}\n"
        f"Price per Kg : {price_per_kg}\n"
        f"Total Price  : {total_price}\n"
        "----------------------------------------\n"
        f"Date/Time    : {current_datetime}\n"
        "----------------------------------------\n"
        "              THANK YOU!                \n"
        "----------------------------------------\n"
    )
    return label_text

# def create_label_text(item_name, weight, price_per_kg, total_price, barcode_data):
#     # Generate barcode image
#     code128 = barcode.get_barcode_class('code128')
#     barcode_img = code128(barcode_data, writer=ImageWriter())

#     # Format the label text
#     label_info = {
#         "Item Name": item_name,
#         "Weight": weight,
#         "Price per Kg": price_per_kg,
#         "Total Price": total_price,
#         "Barcode": barcode_data,
#         "Barcode Image": barcode_img
#     }
#     return label_info



def create_label_file(label_text):
    file_path = tempfile.mktemp(suffix=".txt")
    with open(file_path, "w") as f:
        f.write(label_text)
    return file_path


# def generate_barcode(barcode_text):
#     ean = barcode.get('ean13', barcode_text, writer=ImageWriter())
#     barcode_image = ean.render()
#     return barcode_image

# def create_label_file(label_info):
#     # Create a new Word document
#     doc = Document()

#     # Add text to the document
#     doc.add_heading("ATMART MALL", level=1)
#     doc.add_paragraph("Item Name    : " + label_info["Item Name"])
#     doc.add_paragraph("Weight       : " + str(label_info["Weight"]))  # Convert to string
#     doc.add_paragraph("Price per Kg : " + str(label_info["Price per Kg"]))
#     doc.add_paragraph("Total Price  : " + str(label_info["Total Price"]))
#     doc.add_paragraph("THANK YOU!")

#      # Convert the barcode image data to bytes
#     barcode_image = generate_barcode(label_info["Barcode"])
#     label_info["Barcode Image"] = barcode_image
#     barcode_image_data = label_info["Barcode Image"]
#     img_byte_array = io.BytesIO()
#     barcode_image_data.save(img_byte_array, 'PNG')  # Save the image to the byte array
#     img_byte_array.seek(0)

#     # Save the barcode image as a file that docx can use
#     barcode_image_path = tempfile.mktemp(suffix=".png")
#     with open(barcode_image_path, 'wb') as f:
#         f.write(img_byte_array.getbuffer())

#     # Add the barcode image to the document
#     doc.add_picture(barcode_image_path, width=Inches(2))  # Adjust width as needed

#     # Save the document
#     file_path = tempfile.mktemp(suffix=".docx")
#     doc.save(file_path)

#     return file_path

# def print_file(printer_name, file_path):
#     try:
#         # Open the printer
#         printer_handle = win32print.OpenPrinter(printer_name)
#         # Start the document
#         win32print.StartDocPrinter(printer_handle, 1, ("Label", None, "RAW"))
        
#         # Start the page
#         win32print.StartPagePrinter(printer_handle)
        
#         with open(file_path, "rb") as f:
#             data = f.read()
#             win32print.WritePrinter(printer_handle, data)
            
#         # Send a form feed command to the printer (to ensure auto feed)
#         form_feed = b'\x0C'  # Form feed character in ASCII
#         win32print.WritePrinter(printer_handle, form_feed)
        
#         # End the page
#         win32print.EndPagePrinter(printer_handle)
#         # End the document
#         win32print.EndDocPrinter(printer_handle)
#         # Close the printer
#         win32print.ClosePrinter(printer_handle)
#     except Exception as e:
#         messagebox.showerror("Printing Error", f"Failed to print: {e}")

# Database connection
mydb = mysql.connect(
    host="localhost",
    user="Atmart_DB",
    passwd="login12DB",
    database="foodpurchasinghistory"
)
cursor = mydb.cursor()


sql_insert_or_update_2 = """
UPDATE stock_details
SET Stock = Stock - %s
WHERE Item_Name = %s;
"""
sqlinfo = "INSERT INTO history (item_name, weight, price_per_kg, total_price) VALUES (%s, %s, %s, %s)"
sqlimage = "INSERT INTO image_table (name, image) VALUES (%s, %s)"


 # Example default values and options for dropdowns
default_item_name = "apple"
item_name_options = ["apple", "banana", "cherry", "grapes", "kiwi", "lemon", "mango", "orange", "papaya", "pear", "pineapple", "plum", "pomegranate", "strawberry", "watermelon","corn"]

# Function to run YOLOv8 model and yield annotated frames
def run_yolov8_model():
    # Load YOLOv8 model
    cap = cv2.VideoCapture(0)

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break

        results = model(frame, conf=0.4)
        annotated_frame = results[0].plot()
       
        yield annotated_frame

    cap.release()

# Function to handle "Insert Image" button click
def on_insert_image_button_click():
    name = item_name_var.get()
    # Open camera
    cap = cv2.VideoCapture(0)
    # Check if camera opened successfully
    if not cap.isOpened():
        print("Error: Could not open camera")
        return
    # Capture frame-by-frame
    ret, frame = cap.read()
    
    # Release the camera
    cap.release()
    if not ret:
        print("Error: Could not read frame")
        return
    # Convert image to RGB format
    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    # Store the captured image as a global variable
    captured_image = frame_rgb
    # Convert image to binary
    image_binary = cv2.imencode('.jpg', captured_image)[1].tobytes()
    # Insert the image and name into the database
    cursor.execute(sqlimage, (name, image_binary))
    mydb.commit()
    print("Image inserted successfully")

# Function to start video capture
def start_video():
    global running, video_thread, update_id, auto_update_enabled
    running = True
    update_id = root.after(5000, update_values_item)
    root.after(2000,update_weight)
    auto_update_enabled = not auto_update_enabled
    auto_update_button.config(text="Auto-Update_Stop")
    video_thread = threading.Thread(target=video_loop)
    video_thread.start()

# Function to stop video capture
def stop_video():
    global running
    running = False
    if video_thread and video_thread.is_alive():
        video_thread.join()

# Video loop function to update Tkinter label with YOLOv8 output
def video_loop():

    for frame in run_yolov8_model():
        if not running:
            break

        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        image = Image.fromarray(image)
        image = ImageTk.PhotoImage(image)
        results = model(frame, conf=0.4)
        boxes=results[0].boxes.cpu().numpy()
        # print(boxes.cls)
        # print(boxes)
        if boxes.cls.shape[0]>0:
            class_names2 = []
            seen_classes = set()

            for i in range(boxes.cls.shape[0]):
                class_name = class_names[int(boxes.cls[i])]
                if class_name not in seen_classes:
                    class_names2.append(class_name)
                    seen_classes.add(class_name)

            # print(boxes.cls)
            # annotated_frame = image
            if len(class_names2)>0:
             status_field.config(text=class_names2)
             status_field.text=class_names2
             if len(class_names2)==1:
                 global item_detected
                 item_detected = class_names2[0]
            else:
             status_field.config(text="No object detected")
             status_field.text="No object detected"
             
            if len(class_names2)>1:
             status_field2.config(text="Multiple objects detected")
             status_field2.text="Multiple objects detected"
            else:
             status_field2.config(text="")
             status_field2.text=""
        video_label.config(image=image)
        video_label.image = image



# Placeholder function for next button click
def on_next_button_click():
     # Inserting data into database
    item_name = item_name_combobox.get()
    weight =  Decimal(weight_var.cget("text")).quantize(Decimal('1.000'))
    price_per_kg = Decimal(price_per_kg_var.cget("text")).quantize(Decimal('1.000'))
    total_price = Decimal(total_price_var.cget("text")).quantize(Decimal('1.000'))
    val = (item_name_var.get(), weight, price_per_kg, total_price)

    if item_name and weight:
            try:
                cursor.execute(sql_insert_or_update_2, (weight,item_name))
                cursor.execute(sqlinfo, val)
                mydb.commit()
                print(cursor.rowcount, "record inserted.")
            except mysql.connector.Error as err:
                messagebox.showerror("Error", f"Error updating stock: {err}")
    else:
            messagebox.showwarning("Input Error", "Please enter both stock name and quantity")


    
    label_text = create_label_text(item_name, str(weight), str(price_per_kg), str(total_price))
    file_path = create_label_file(label_text)
    # print_file("POS80 Printer", file_path)

 # Function to create an entry field with a dropdown
def create_entry_with_dropdown(frame, text, row, default_value, options, width=20):
    label = tk.Label(frame, text=text, font=('Helvetica', 12, 'bold'), bg='white', fg='black')
    label.grid(row=row, column=0, padx=5, pady=5, sticky=tk.E)

    # Create a StringVar to hold the current value
    current_value = tk.StringVar(value=default_value)

    # Create the dropdown (combobox)
    combobox = ttk.Combobox(frame, textvariable=current_value, values=options, font=('Helvetica', 12, 'bold'), width=width)
    combobox.grid(row=row, column=1, padx=5, pady=5)
    combobox.set(default_value)

    return combobox, current_value
# Function to update combobox values with detected values
# Placeholder function for object detection
def detection_function():
    return {
        "item_name": item_detected,
        "price_per_kg": Decimal(price[item_detected]).quantize(Decimal('1.000'))
    }

def update_values_item():
    new_values = detection_function()
    item_name_combobox.set(new_values["item_name"])
    item_name_var.set(new_values["item_name"])
    price_per_kg_var.config(text=new_values["price_per_kg"])

    global update_id
    # Schedule the next update
    if(auto_update_enabled):
        update_id = root.after(5000, update_values_item)
        

def update_weight():
    line = ser.readline().decode('utf-8').strip()
    if line:
        try:
            global weight
            weight = round((float(line)/1000)+0.0001,3)
            weight_var.config(text=str(weight))
            total_price_var.config(text=str(weight*int(price_per_kg_var.cget("text"))))
                # w = weight  # Store the weight in the variable w
            print(f"Weight: {weight} g")
        except ValueError:
                # Handle the case where the serial data is not a valid float
            print(f"Received invalid data: {line}")
    
    root.after(500, update_weight)
# Function to quit the application
def on_quit_button_click():
    check_password()

def on_closing():
    check_password()

def check_password():
    password = simpledialog.askstring("Password", "Enter password:", show='*')
    if password == "1234":  # Replace "your_password" with the correct password
            # ser.close()
            mydb.close()
            root.quit()
            root.destroy()
    else:
        messagebox.showerror("Error", "Wrong password")
        # Automatically close the error message box after 5 seconds
        root.after(5000, lambda: messagebox._show)

auto_update_enabled = False  # or True, depending on your initial state

def toggle_auto_update():
    global auto_update_enabled
    global update_id
    if auto_update_enabled:
        auto_update_enabled = not auto_update_enabled
        auto_update_button.config(text="Auto-Update_Start")
        
        root.after_cancel(update_id)
    else:
        update_id = root.after(5000, update_values_item)
        root.after(2000,update_weight)
        auto_update_enabled = not auto_update_enabled
        auto_update_button.config(text="Auto-Update_Stop")


# Placeholder function for login
def login():
    global login_message
    username = username_entry.get()
    password = password_entry.get()
    if username == "" and password == "":
        login_window.destroy()
        show_dashboard()
    else:
        login_message.config(text="Please check your username or password")

# Function to display the main dashboard window
def show_dashboard():
    global video_label, root, running,status_field,status_field2,item_name_combobox, item_name_var,weight_var, price_per_kg_var, total_price_var, update_button,name_entry,auto_update_button

    root = tk.Tk()
    root.title('Vegetables and Fruits Recognition Based Auto-Packaging Management Dashboard')
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.geometry(f"{root.winfo_screenwidth()}x{root.winfo_screenheight()}")
    root.configure(background='white')

    img_path = r'/home/asus/Desktop/Interface/guienv/lib/python3.11/site-packages/Logo_of_IIT_Bhilai.png'
    bg_image = Image.open(img_path)
    resized_bg_image = bg_image.resize((100, 100), Image.Resampling.LANCZOS)
    bg_image = ImageTk.PhotoImage(resized_bg_image)
    root.bg_image = bg_image

    header_frame = tk.Frame(root, bg='white')
    header_frame.pack(fill=tk.X, pady=10)

    img_label_left = tk.Label(header_frame, image=root.bg_image, bg='white')
    img_label_left.pack(side=tk.LEFT, padx=20)

    img_label_right = tk.Label(header_frame, image=root.bg_image, bg='white')
    img_label_right.pack(side=tk.RIGHT, padx=20)

    title_label = tk.Label(header_frame, text="VEGETABLES AND FRUITS RECOGNITION BASED AUTO-PACKAGING MANAGEMENT DASHBOARD", font=('Helvetica', 18, 'bold'), bg='white', fg='black', pady=20)
    title_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

    main_frame = tk.Frame(root, bg='white')
    main_frame.pack(pady=20)
    

    # Create a single frame for the action
    action_frame = tk.LabelFrame(main_frame, text="Action", font=('Helvetica', 14, 'bold'), bg='white', fg='black', padx=20, pady=20, width=500, height=300)
    action_frame.grid(row=0, column=0, padx=20, pady=10)
    action_frame.grid_propagate(False)  # Prevent the frame from resizing to fit its contents

    def create_display_field(frame, text, row, default_value):
        label = tk.Label(frame, text=text, font=('Helvetica', 12, 'bold'), bg='white', fg='black')
        label.grid(row=row, column=0, padx=5, pady=5, sticky=tk.E)
        value_label = tk.Label(frame, text=default_value, font=('Helvetica', 12, 'bold'), bg='white', fg='black')
        value_label.grid(row=row, column=1, padx=5, pady=5, sticky=tk.W)
        return value_label

    # Create entry fields with dropdowns
    item_name_combobox, item_name_var = create_entry_with_dropdown(action_frame, "Item Name", 0, default_item_name, item_name_options, width=30)
    weight_var= create_display_field(action_frame, "Weight(Kg)", 1, "2")
    price_per_kg_var = create_display_field(action_frame, "Price/Kg(Rs.)", 2, "0")
    total_price_var = create_display_field(action_frame, "Total Price(Rs.)", 3, "0")

    update_button = tk.Button(action_frame, text="Update values", font=('Helvetica', 14, 'bold'), bg='#2196F3', fg='white', padx=10, pady=5, command=update_values_item)
    update_button.grid(row=4, column=0, pady=10)

    auto_update_button = tk.Button(action_frame, text="Auto_update_Start", font=('Helvetica', 14, 'bold'), bg='#2196F3', fg='white', padx=10, pady=5, command=toggle_auto_update)
    auto_update_button.grid(row=4, column=1, pady=10)

    live_recognition_frame = tk.LabelFrame(main_frame, text="Live Recognition", font=('Helvetica', 14, 'bold'), bg='white', fg='black', padx=20, pady=20)
    live_recognition_frame.grid(row=0, column=2, padx=20, pady=10)

    placeholder_img = Image.new('RGBA', (540, 440), (200, 200, 200, 255))
    placeholder_image = ImageTk.PhotoImage(placeholder_img)

    video_label = tk.Label(live_recognition_frame, image=placeholder_image, width=540, height=440, bg='white')
    video_label.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

    start_button = tk.Button(live_recognition_frame, text="Start", font=('Helvetica', 12, 'bold'), bg='#4CAF50', fg='white', padx=10, command=start_video)
    start_button.grid(row=1, column=0, padx=5, pady=10)
    
    stop_button = tk.Button(live_recognition_frame, text="Stop", font=('Helvetica', 12, 'bold'), bg='#f44336', fg='white', padx=10, command=stop_video)
    stop_button.grid(row=1, column=1, padx=5, pady=10)


    status_frame = tk.LabelFrame(main_frame, text="Current Status", font=('Helvetica', 14, 'bold'), bg='white', fg='black', padx=20, pady=20)
    status_frame.grid(row=1,padx=5, pady=10)

    status_field = tk.Label(status_frame, text="No object detected", font=('Helvetica', 12, 'bold'), bg='white', fg='black', width=30)
    status_field.grid(row=1, column=0, padx=10, pady=10)
    status_field2 = tk.Label(status_frame, text="", font=('Helvetica', 14, 'bold'), bg='white', fg='red', width=30)
    status_field2.grid(row=1, column=1, padx=10, pady=10)

    button_frame = tk.Frame(root, bg='white')
    button_frame.pack(pady=20)
    quit_button = tk.Button(button_frame, text="Quit", font=('Helvetica', 14, 'bold'), bg='#f44336', fg='white', padx=10, pady=5, command=on_quit_button_click)
    quit_button.pack(side=tk.LEFT, padx=10)
    insert_button = tk.Button(button_frame, text="Insert Image", font=('Helvetica', 14, 'bold'), bg='#2196F3', fg='white', padx=5, command=on_insert_image_button_click)
    insert_button.pack(side=tk.RIGHT, padx=10)
    next_button = tk.Button(button_frame, text="Next", font=('Helvetica', 14, 'bold'), bg='#2196F3', fg='white', padx=10, pady=5, command=on_next_button_click)
    next_button.pack(side=tk.RIGHT, padx=10)
  
    

    root.mainloop()

# Create login window
login_window = tk.Tk()
login_window.title("Vegetables and Fruits Recognition Based Auto-Packaging Management Dashboard")
login_window.geometry(f"{login_window.winfo_screenwidth()}x{login_window.winfo_screenheight()}")
login_window.configure(background='white')

# Title Label
login_title_label = tk.Label(login_window, text="VEGETABLES AND FRUITS RECOGNITION BASED AUTO-PACKAGING MANAGEMENT DASHBOARD", font=('Helvetica', 18, 'bold'), bg='white', fg='black')
login_title_label.place(relx=0.5, rely=0.2, anchor=tk.CENTER)

# Username Label and Entry
username_label = tk.Label(login_window, text="Username:", font=('Helvetica', 16, 'bold'), bg='white', fg='black')
username_label.place(relx=0.5, rely=0.4, anchor=tk.CENTER)
username_entry = tk.Entry(login_window, font=('Helvetica', 14, 'bold'))
username_entry.place(relx=0.5, rely=0.45, anchor=tk.CENTER)

# Password Label and Entry
password_label = tk.Label(login_window, text="Password:", font=('Helvetica', 16, 'bold'), bg='white', fg='black')
password_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
password_entry = tk.Entry(login_window, show="*", font=('Helvetica', 14, 'bold'))
password_entry.place(relx=0.5, rely=0.55, anchor=tk.CENTER)

# Login Button
login_button = tk.Button(login_window, text="Login", font=('Helvetica', 14, 'bold'), command=login)
login_button.place(relx=0.5, rely=0.65, anchor=tk.CENTER)

#  Login Message
login_message = tk.Label(login_window, text="", font=('Helvetica', 12, 'bold'), bg='white', fg='red')
login_message.place(relx=0.5, rely=0.75, anchor=tk.CENTER)

# Start the Tkinter main loop
login_window.mainloop()

